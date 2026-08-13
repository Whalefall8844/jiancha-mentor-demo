from __future__ import annotations

from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document

from ..repositories.catalog import (
    get_rule_pack,
    get_template,
    list_template_mappings,
    create_configuration_audit_event,
    update_rule_pack_control,
    update_template_control,
)
from .system_checks import normalize_system_checks
from .escalation_sla import validate_escalation_sla_configuration


EDITABLE_STATUSES = {"draft", "rejected"}


def _now() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M")


def _parse_date(value: str, field_label: str, errors: list[str]) -> datetime | None:
    if not value:
        errors.append(f"{field_label}不能为空")
        return None
    try:
        return datetime.strptime(value, "%Y-%m-%d")
    except ValueError:
        errors.append(f"{field_label}应使用 YYYY-MM-DD 格式")
        return None


def validate_template_activation(template_id: str) -> dict[str, Any]:
    template = get_template(template_id)
    if template is None:
        raise ValueError("未找到 Word 模板")

    errors: list[str] = []
    document: Document | None = None
    template_path = Path(template.get("docx_path") or "")
    if not template_path.exists():
        errors.append("模板原文件不存在或不可读取")
    else:
        try:
            document = Document(template_path)
        except Exception:
            errors.append("模板原文件无法打开")

    actual_table_count = len(document.tables) if document is not None else 0
    expected_table_count = int(template.get("table_count") or 0)
    if expected_table_count <= 0:
        errors.append("模板未识别到可配置表格")
    elif actual_table_count != expected_table_count:
        errors.append(f"模板表格数量变化：登记 {expected_table_count} 张，当前读取 {actual_table_count} 张")

    mappings = list_template_mappings(template_id)
    indexes = [int(item["table_index"]) for item in mappings]
    expected_indexes = set(range(1, expected_table_count + 1))
    actual_indexes = set(indexes)
    if actual_indexes != expected_indexes or len(indexes) != len(actual_indexes):
        errors.append("字段映射未覆盖每一张表，或存在重复表格映射")
    if any(not str(item.get("target_description") or "").strip() for item in mappings):
        errors.append("存在未命名的监查任务/报告区域")
    field_keys = [str(item.get("field_key") or "").strip() for item in mappings]
    if len(field_keys) != len(set(field_keys)) or any(not key for key in field_keys):
        errors.append("字段键存在重复或为空，无法形成唯一映射")

    round_trip_table_count = 0
    if document is not None:
        try:
            output = BytesIO()
            document.save(output)
            output.seek(0)
            round_trip_table_count = len(Document(output).tables)
            if round_trip_table_count != actual_table_count:
                errors.append("模板测试生成后表格结构不一致")
        except Exception:
            errors.append("模板测试生成失败")

    return {
        "passed": not errors,
        "checked_at": _now(),
        "errors": errors,
        "registered_table_count": expected_table_count,
        "read_table_count": actual_table_count,
        "test_generated_table_count": round_trip_table_count,
        "mapping_count": len(mappings),
    }


def validate_rule_pack_activation(rule_pack_id: str) -> dict[str, Any]:
    rule_pack = get_rule_pack(rule_pack_id)
    if rule_pack is None:
        raise ValueError("未找到规则包")

    errors: list[str] = []
    if not str(rule_pack.get("name") or "").strip():
        errors.append("规则包名称不能为空")
    if not str(rule_pack.get("version") or "").strip():
        errors.append("规则包版本不能为空")
    effective_from = _parse_date(str(rule_pack.get("effective_from") or ""), "生效日期", errors)
    effective_to_value = str(rule_pack.get("effective_to") or "").strip()
    effective_to = None
    if effective_to_value:
        try:
            effective_to = datetime.strptime(effective_to_value, "%Y-%m-%d")
        except ValueError:
            errors.append("失效日期应使用 YYYY-MM-DD 格式")
    if effective_from and effective_to and effective_to < effective_from:
        errors.append("失效日期不得早于生效日期")
    content = rule_pack.get("content")
    if not isinstance(content, dict) or not content:
        errors.append("规则包内容不能为空")
        system_checks: list[dict[str, Any]] = []
    else:
        try:
            system_checks = normalize_system_checks(content)
        except ValueError as exc:
            errors.append(str(exc))
            system_checks = []
        errors.extend(validate_escalation_sla_configuration(content))

    return {
        "passed": not errors,
        "checked_at": _now(),
        "errors": errors,
        "effective_from": rule_pack.get("effective_from", ""),
        "effective_to": rule_pack.get("effective_to", ""),
        "content_keys": sorted(str(key) for key in (content or {}).keys()),
        "system_check_count": len(system_checks),
    }


def _template_metadata_with_validation(template: dict[str, Any], validation: dict[str, Any]) -> dict[str, Any]:
    metadata = dict(template.get("metadata") or {})
    metadata["activation_check"] = validation
    return metadata


def transition_template(*, template_id: str, action: str, actor_name: str, note: str = "") -> dict[str, Any]:
    before = get_template(template_id)
    if before is None:
        raise ValueError("未找到 Word 模板")
    status = before.get("status", "draft")
    actor = actor_name.strip() or "演示审批人"
    review_note = note.strip()
    validation: dict[str, Any] | None = None

    if action == "submit":
        if status not in EDITABLE_STATUSES:
            raise ValueError("只有草稿或已退回模板可以提交审批")
        validation = validate_template_activation(template_id)
        if not validation["passed"]:
            raise ValueError("模板未通过激活检查：" + "；".join(validation["errors"]))
        updated = update_template_control(
            template_id,
            {
                "status": "pending_approval",
                "submitted_at": _now(),
                "submitted_by": actor,
                "reviewed_at": "",
                "reviewed_by": "",
                "review_note": "",
                "metadata": _template_metadata_with_validation(before, validation),
            },
        )
    elif action == "approve":
        if status != "pending_approval":
            raise ValueError("只有待审批模板可以批准启用")
        validation = validate_template_activation(template_id)
        if not validation["passed"]:
            raise ValueError("模板当前未通过激活检查：" + "；".join(validation["errors"]))
        updated = update_template_control(
            template_id,
            {
                "status": "active",
                "reviewed_at": _now(),
                "reviewed_by": actor,
                "review_note": review_note,
                "metadata": _template_metadata_with_validation(before, validation),
            },
        )
    elif action == "reject":
        if status != "pending_approval":
            raise ValueError("只有待审批模板可以退回")
        if not review_note:
            raise ValueError("退回模板时必须填写审批意见")
        updated = update_template_control(
            template_id,
            {"status": "rejected", "reviewed_at": _now(), "reviewed_by": actor, "review_note": review_note},
        )
    elif action == "withdraw":
        if status != "pending_approval":
            raise ValueError("只有待审批模板可以撤回")
        updated = update_template_control(template_id, {"status": "draft", "review_note": review_note})
    elif action == "deactivate":
        if status != "active":
            raise ValueError("只有已启用模板可以停用")
        updated = update_template_control(
            template_id,
            {"status": "inactive", "reviewed_at": _now(), "reviewed_by": actor, "review_note": review_note},
        )
    else:
        raise ValueError("不支持的模板审批操作")
    item = updated or {}
    create_configuration_audit_event(
        entity_type="template",
        entity_id=template_id,
        action=f"approval_{action}",
        actor_name=actor,
        detail={"from_status": status, "to_status": item.get("status", ""), "note": review_note, "validation": validation},
    )
    return {"item": item, "before_status": status, "validation": validation}


def transition_rule_pack(*, rule_pack_id: str, action: str, actor_name: str, note: str = "") -> dict[str, Any]:
    before = get_rule_pack(rule_pack_id)
    if before is None:
        raise ValueError("未找到规则包")
    status = before.get("status", "draft")
    actor = actor_name.strip() or "演示审批人"
    review_note = note.strip()
    validation: dict[str, Any] | None = None

    if action == "submit":
        if status not in EDITABLE_STATUSES:
            raise ValueError("只有草稿或已退回规则包可以提交审批")
        validation = validate_rule_pack_activation(rule_pack_id)
        if not validation["passed"]:
            raise ValueError("规则包未通过启用检查：" + "；".join(validation["errors"]))
        updated = update_rule_pack_control(
            rule_pack_id,
            {
                "status": "pending_approval",
                "submitted_at": _now(),
                "submitted_by": actor,
                "reviewed_at": "",
                "reviewed_by": "",
                "review_note": "",
            },
        )
    elif action == "approve":
        if status != "pending_approval":
            raise ValueError("只有待审批规则包可以批准启用")
        validation = validate_rule_pack_activation(rule_pack_id)
        if not validation["passed"]:
            raise ValueError("规则包当前未通过启用检查：" + "；".join(validation["errors"]))
        updated = update_rule_pack_control(
            rule_pack_id,
            {"status": "active", "reviewed_at": _now(), "reviewed_by": actor, "review_note": review_note},
        )
    elif action == "reject":
        if status != "pending_approval":
            raise ValueError("只有待审批规则包可以退回")
        if not review_note:
            raise ValueError("退回规则包时必须填写审批意见")
        updated = update_rule_pack_control(
            rule_pack_id,
            {"status": "rejected", "reviewed_at": _now(), "reviewed_by": actor, "review_note": review_note},
        )
    elif action == "withdraw":
        if status != "pending_approval":
            raise ValueError("只有待审批规则包可以撤回")
        updated = update_rule_pack_control(rule_pack_id, {"status": "draft", "review_note": review_note})
    elif action == "deactivate":
        if status != "active":
            raise ValueError("只有已启用规则包可以停用")
        updated = update_rule_pack_control(
            rule_pack_id,
            {"status": "inactive", "reviewed_at": _now(), "reviewed_by": actor, "review_note": review_note},
        )
    else:
        raise ValueError("不支持的规则包审批操作")
    item = updated or {}
    create_configuration_audit_event(
        entity_type="rule_pack",
        entity_id=rule_pack_id,
        project_id=str(before.get("project_id") or ""),
        action=f"approval_{action}",
        actor_name=actor,
        detail={"from_status": status, "to_status": item.get("status", ""), "note": review_note, "validation": validation},
    )
    return {"item": item, "before_status": status, "validation": validation}
