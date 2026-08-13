from __future__ import annotations

import re
import unicodedata
from pathlib import Path
from typing import Any

from docx import Document

from .template_slots import find_inline_tokens


FIELD_SLOT_SUGGESTION_ALGORITHM = "template_field_slot_label_v1"


SLOT_SOURCE_PROFILES: tuple[dict[str, Any], ...] = (
    {"value_source": "project.study_name", "field_key": "project_study_name", "label": "项目名称", "terms": ("项目名称", "研究名称", "试验名称", "study name")},
    {"value_source": "project.study_id", "field_key": "project_study_id", "label": "项目/方案编号", "terms": ("项目编号", "研究编号", "方案编号", "试验编号", "study no", "study id")},
    {"value_source": "project.sponsor", "field_key": "project_sponsor", "label": "申办方", "terms": ("申办方", "申办者", "sponsor")},
    {"value_source": "project.approval_number", "field_key": "project_approval_number", "label": "立项/批件号", "terms": ("批准文号", "批件号", "立项号", "批号", "approval no")},
    {"value_source": "project.sop_version", "field_key": "project_sop_version", "label": "监查 SOP 版本", "terms": ("sop版本", "监查sop", "sop version")},
    {"value_source": "site.site_name", "field_key": "site_name", "label": "中心名称", "terms": ("中心名称", "研究中心", "医院名称", "site name")},
    {"value_source": "site.pi_name", "field_key": "site_pi_name", "label": "主要研究者（PI）", "terms": ("主要研究者", "中心pi", "研究者", "principal investigator", "pi")},
    {"value_source": "site.protocol_version", "field_key": "site_protocol_version", "label": "方案版本", "terms": ("方案版本", "方案号", "protocol version", "protocol")},
    {"value_source": "site.icf_version", "field_key": "site_icf_version", "label": "知情同意书版本", "terms": ("知情同意书版本", "知情同意版本", "icf版本", "icf version")},
    {"value_source": "site.ethics_date", "field_key": "site_ethics_date", "label": "伦理日期", "terms": ("伦理批准日期", "伦理日期", "伦理批件日期", "ec date", "irb date")},
    {"value_source": "visit.activity_period", "field_key": "visit_activity_period", "label": "监查活动周期", "terms": ("监查日期", "访视日期", "监查时间", "监查期间", "visit date")},
    {"value_source": "visit.visit_method", "field_key": "visit_method", "label": "监查方式", "terms": ("监查方式", "访视方式", "监查类型", "visit method")},
    {"value_source": "visit.report_date", "field_key": "visit_report_date", "label": "报告日期", "terms": ("报告日期", "报告完成日期", "report date")},
    {"value_source": "visit.site_team", "field_key": "visit_site_team", "label": "中心研究团队", "terms": ("中心研究团队", "研究团队", "中心人员", "研究人员")},
    {"value_source": "visit.monitoring_team", "field_key": "visit_monitoring_team", "label": "监查团队", "terms": ("监查人员", "监查员", "cra", "监查团队")},
    {"value_source": "visit.next_visit", "field_key": "visit_next_visit", "label": "下次访视计划", "terms": ("下次访视计划", "下次监查计划", "下次访视", "后续计划")},
    {"value_source": "summary", "field_key": "report_summary", "label": "本次监查总体评价", "terms": ("总体评价", "总体结论", "监查结论", "监查总结", "总结与建议")},
)


def _normalise(value: object) -> str:
    text = unicodedata.normalize("NFKC", str(value or "")).casefold()
    return re.sub(r"[^\w\u4e00-\u9fff]+", "", text)


def _compact(value: object) -> str:
    return " ".join(str(value or "").split()).strip()


def _matching_profile(value: str) -> tuple[dict[str, Any] | None, list[str]]:
    normalized = _normalise(value)
    best_profile: dict[str, Any] | None = None
    best_terms: list[str] = []
    best_score = 0
    for profile in SLOT_SOURCE_PROFILES:
        matched_terms = [term for term in profile["terms"] if _normalise(term) and _normalise(term) in normalized]
        score = sum(len(_normalise(term)) for term in matched_terms)
        if score > best_score:
            best_profile = profile
            best_terms = matched_terms
            best_score = score
    return best_profile, list(dict.fromkeys(best_terms))


def _blank_neighbor(table, source_row: int, source_column: int) -> tuple[int, int] | None:
    source_cell = table.cell(source_row, source_column)
    source_identity = id(source_cell._tc)
    for column_index in range(source_column + 1, len(table.columns)):
        candidate = table.cell(source_row, column_index)
        if id(candidate._tc) != source_identity and not _compact(candidate.text):
            return source_row, column_index
    for row_index in range(source_row + 1, len(table.rows)):
        candidate = table.cell(row_index, source_column)
        if id(candidate._tc) != source_identity and not _compact(candidate.text):
            return row_index, source_column
    return None


def _suggestion_confidence(matched_terms: list[str]) -> str:
    return "high" if max((len(_normalise(term)) for term in matched_terms), default=0) >= 5 else "medium"


def _unique_field_key(base_key: str, used_field_keys: dict[str, int]) -> tuple[str, int]:
    key_number = used_field_keys.get(base_key, 0) + 1
    used_field_keys[base_key] = key_number
    return (base_key if key_number == 1 else f"{base_key}_{key_number}", key_number)


def _inline_token_suggestions(
    document: Document,
    *,
    used_targets: set[str],
    used_field_keys: dict[str, int],
) -> list[dict[str, Any]]:
    suggestions: list[dict[str, Any]] = []

    def add_for_paragraph(*, target_prefix: str, paragraph_text: str, location_label: str) -> None:
        for token in find_inline_tokens(paragraph_text):
            token_label = _compact(token.removeprefix("{{").removesuffix("}}"))
            profile, matched_terms = _matching_profile(token_label)
            if profile is None:
                continue
            target_locator = f"{target_prefix}:{token}"
            if target_locator in used_targets:
                continue
            used_targets.add(target_locator)
            field_key, key_number = _unique_field_key(str(profile["field_key"]), used_field_keys)
            reason = f"{location_label}内联标记“{token}”命中：{'、'.join(matched_terms)}。"
            if key_number > 1:
                reason += " 同类标记重复出现，字段键增加序号以保持独立。"
            suggestions.append({
                "table_index": 0,
                "target_kind": "inline_token",
                "target_locator": target_locator,
                "label": str(profile["label"]),
                "field_key": field_key,
                "value_source": str(profile["value_source"]),
                "confidence": _suggestion_confidence(matched_terms),
                "matched_terms": matched_terms,
                "reason": reason,
                "algorithm": FIELD_SLOT_SUGGESTION_ALGORITHM,
            })

    for paragraph_index, paragraph in enumerate(document.paragraphs, start=1):
        add_for_paragraph(
            target_prefix=f"P{paragraph_index}",
            paragraph_text=paragraph.text,
            location_label=f"正文第 {paragraph_index} 段",
        )
    for section_index, section in enumerate(document.sections, start=1):
        for paragraph_index, paragraph in enumerate(section.header.paragraphs, start=1):
            add_for_paragraph(
                target_prefix=f"H{section_index}:P{paragraph_index}",
                paragraph_text=paragraph.text,
                location_label=f"第 {section_index} 节页眉第 {paragraph_index} 段",
            )
        for paragraph_index, paragraph in enumerate(section.footer.paragraphs, start=1):
            add_for_paragraph(
                target_prefix=f"F{section_index}:P{paragraph_index}",
                paragraph_text=paragraph.text,
                location_label=f"第 {section_index} 节页脚第 {paragraph_index} 段",
            )
    return suggestions


def suggest_template_field_slots(docx_path: Path | str) -> list[dict[str, Any]]:
    """Infer fixed-data field slots from table labels and Word inline markers."""
    document = Document(Path(docx_path))
    suggestions: list[dict[str, Any]] = []
    used_targets: set[str] = set()
    used_field_keys: dict[str, int] = {}
    for table_index, table in enumerate(document.tables, start=1):
        seen_source_cells: set[int] = set()
        for row_index, row in enumerate(table.rows):
            for column_index, cell in enumerate(row.cells):
                cell_identity = id(cell._tc)
                if cell_identity in seen_source_cells:
                    continue
                seen_source_cells.add(cell_identity)
                label_text = _compact(cell.text)
                if not label_text:
                    continue
                profile, matched_terms = _matching_profile(label_text)
                if profile is None:
                    continue
                neighbor = _blank_neighbor(table, row_index, column_index)
                if neighbor is None:
                    continue
                target_row, target_column = neighbor
                target_locator = f"T{table_index}:R{target_row + 1}:C{target_column + 1}"
                if target_locator in used_targets:
                    continue
                used_targets.add(target_locator)
                base_key = str(profile["field_key"])
                field_key, key_number = _unique_field_key(base_key, used_field_keys)
                confidence = _suggestion_confidence(matched_terms)
                location = "右侧" if target_row == row_index else "下方"
                reason = f"标签“{label_text[:80]}”命中：{'、'.join(matched_terms)}；建议写入其{location}空白单元格。"
                if key_number > 1:
                    reason += " 同类标签重复出现，字段键增加序号以保持独立。"
                suggestions.append({
                    "table_index": table_index,
                    "target_kind": "table_cell",
                    "target_locator": target_locator,
                    "label": str(profile["label"]),
                    "field_key": field_key,
                    "value_source": str(profile["value_source"]),
                    "confidence": confidence,
                    "matched_terms": matched_terms,
                    "reason": reason,
                    "algorithm": FIELD_SLOT_SUGGESTION_ALGORITHM,
                })
    suggestions.extend(
        _inline_token_suggestions(
            document,
            used_targets=used_targets,
            used_field_keys=used_field_keys,
        )
    )
    return suggestions
