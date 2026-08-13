from __future__ import annotations

from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any
from uuid import uuid4

from docx import Document

from ..database import BACKEND_DIR
from ..repositories.catalog import (
    create_template,
    create_template_field_slot,
    delete_template_field_slot,
    get_template,
    list_template_field_slots,
    list_templates,
    list_template_mappings,
    replace_template_field_slots,
    replace_template_mappings,
    update_template_control,
    update_template_document,
    update_template_field_slot,
    update_template_mapping,
)
from .template_matching import build_matching_profile, normalize_visit_type_keywords, recommend_templates
from .template_mapping_suggestions import suggest_template_mappings
from .template_field_slot_suggestions import suggest_template_field_slots
from .docx_structured_targets import discover_structured_targets
from .template_slots import find_inline_tokens, parse_slot_target


TEMPLATE_UPLOAD_DIR = BACKEND_DIR / "uploads" / "templates"


def _table_label(table, table_index: int) -> str:
    """Return the first usable heading-like cell so an administrator can refine it later."""
    for row in list(table.rows)[:2]:
        for cell in list(row.cells)[:4]:
            value = " ".join(cell.text.split())
            if value:
                return value[:160]
    return f"表 {table_index}"


def _first_blank_cell_locator(table, table_index: int) -> str:
    seen_cells: set[int] = set()
    for row_index, row in enumerate(table.rows, start=1):
        for column_index, cell in enumerate(row.cells, start=1):
            cell_identity = id(cell._tc)
            if cell_identity in seen_cells:
                continue
            seen_cells.add(cell_identity)
            if not " ".join(cell.text.split()):
                return f"T{table_index}:R{row_index}:C{column_index}"
    return ""


def _detect_tables(document: Document) -> list[dict[str, Any]]:
    detected_tables: list[dict[str, Any]] = []
    for table_index, table in enumerate(document.tables, start=1):
        detected_tables.append(
            {
                "table_index": table_index,
                "detected_label": _table_label(table, table_index),
                "row_count": len(table.rows),
                "column_count": len(table.columns),
                "suggested_target_locator": _first_blank_cell_locator(table, table_index),
            }
        )
    return detected_tables


def _table_label_signature(value: object) -> str:
    return " ".join(str(value or "").split()).casefold()


def _describe_template_table_changes(
    previous_tables: list[dict[str, Any]],
    current_tables: list[dict[str, Any]],
) -> dict[str, Any]:
    previous_by_index = {int(table.get("table_index") or 0): table for table in previous_tables}
    current_by_index = {int(table.get("table_index") or 0): table for table in current_tables}
    changes: list[dict[str, Any]] = []
    for table_index in sorted(set(previous_by_index) | set(current_by_index)):
        previous = previous_by_index.get(table_index)
        current = current_by_index.get(table_index)
        if previous is None:
            status = "added"
        elif current is None:
            status = "removed"
        elif (
            _table_label_signature(previous.get("detected_label")) == _table_label_signature(current.get("detected_label"))
            and int(previous.get("row_count") or 0) == int(current.get("row_count") or 0)
            and int(previous.get("column_count") or 0) == int(current.get("column_count") or 0)
        ):
            status = "unchanged"
        else:
            status = "changed"
        changes.append(
            {
                "table_index": table_index,
                "status": status,
                "previous_label": str((previous or {}).get("detected_label") or ""),
                "current_label": str((current or {}).get("detected_label") or ""),
                "previous_row_count": int((previous or {}).get("row_count") or 0),
                "previous_column_count": int((previous or {}).get("column_count") or 0),
                "current_row_count": int((current or {}).get("row_count") or 0),
                "current_column_count": int((current or {}).get("column_count") or 0),
            }
        )
    return {
        "unchanged_table_count": sum(1 for change in changes if change["status"] == "unchanged"),
        "changed_table_count": sum(1 for change in changes if change["status"] == "changed"),
        "added_table_count": sum(1 for change in changes if change["status"] == "added"),
        "removed_table_count": sum(1 for change in changes if change["status"] == "removed"),
        "table_changes": changes,
    }


def _text_preview(value: str) -> str:
    text = " ".join(str(value or "").split())
    return text if len(text) <= 160 else text[:159] + "…"


def _describe_text_target(
    *,
    target_kind: str,
    target_locator: str,
    label: str,
    text: str,
) -> list[dict[str, Any]]:
    preview = _text_preview(text)
    if not preview:
        return []
    targets = [
        {
            "target_kind": target_kind,
            "target_locator": target_locator,
            "label": label,
            "preview": preview,
        }
    ]
    for token in find_inline_tokens(text):
        targets.append(
            {
                "target_kind": "inline_token",
                "target_locator": f"{target_locator}:{token}",
                "label": f"{label} 内联标记 {token}",
                "preview": preview,
            }
        )
    return targets


def _detect_text_targets(document: Document, docx_path: Path | None = None) -> list[dict[str, Any]]:
    targets: list[dict[str, Any]] = []
    for paragraph_index, paragraph in enumerate(document.paragraphs, start=1):
        targets.extend(
            _describe_text_target(
                target_kind="body_paragraph",
                target_locator=f"P{paragraph_index}",
                label=f"正文第 {paragraph_index} 段",
                text=paragraph.text,
            )
        )
    for section_index, section in enumerate(document.sections, start=1):
        for paragraph_index, paragraph in enumerate(section.header.paragraphs, start=1):
            targets.extend(
                _describe_text_target(
                    target_kind="header_paragraph",
                    target_locator=f"H{section_index}:P{paragraph_index}",
                    label=f"第 {section_index} 节页眉第 {paragraph_index} 段",
                    text=paragraph.text,
                )
            )
        for paragraph_index, paragraph in enumerate(section.footer.paragraphs, start=1):
            targets.extend(
                _describe_text_target(
                    target_kind="footer_paragraph",
                    target_locator=f"F{section_index}:P{paragraph_index}",
                    label=f"第 {section_index} 节页脚第 {paragraph_index} 段",
                    text=paragraph.text,
                )
            )
    if docx_path is not None:
        targets.extend(discover_structured_targets(docx_path))
    return targets


def _detected_tables_for_template(template: dict[str, Any]) -> list[dict[str, Any]]:
    metadata = template.get("metadata", {})
    detected_tables = metadata.get("detected_tables", [])
    if detected_tables:
        return detected_tables
    template_path = Path(template["docx_path"])
    try:
        return _detect_tables(Document(template_path)) if template_path.exists() else []
    except Exception:
        return []


def _detected_text_targets_for_template(template: dict[str, Any]) -> list[dict[str, Any]]:
    metadata = template.get("metadata", {})
    detected_text_targets = metadata.get("detected_text_targets", [])
    if detected_text_targets:
        return detected_text_targets
    template_path = Path(template["docx_path"])
    try:
        return _detect_text_targets(Document(template_path), template_path) if template_path.exists() else []
    except Exception:
        return []


def _is_unconfigured_default_field_slot(slot: dict[str, Any]) -> bool:
    table_index = int(slot.get("table_index") or 0)
    return (
        str(slot.get("target_kind") or "table_cell") == "table_cell"
        and str(slot.get("value_source") or "") == "confirmed_text"
        and str(slot.get("field_key") or "") == f"table_{table_index}"
        and str(slot.get("label") or "").endswith("填写位")
        and not bool(slot.get("required"))
    )


def _is_unconfigured_default_mapping(mapping: dict[str, Any], detected_table: dict[str, Any]) -> bool:
    table_index = int(mapping.get("table_index") or 0)
    return (
        str(mapping.get("field_key") or "") == f"table_{table_index}"
        and str(mapping.get("target_description") or "").strip()
        == str(detected_table.get("detected_label") or "").strip()
        and not bool(mapping.get("required"))
    )


def _template_configuration_readiness(
    *,
    detected_tables: list[dict[str, Any]],
    mappings: list[dict[str, Any]],
    field_slots: list[dict[str, Any]],
    mapping_suggestions: list[dict[str, Any]],
    field_slot_suggestions: list[dict[str, Any]],
) -> dict[str, Any]:
    """Summarize configuration progress without changing approval or report gates."""
    mappings_by_table = {
        int(mapping.get("table_index") or 0): mapping
        for mapping in mappings
    }
    configured_mapping_count = 0
    pending_mapping_count = 0
    for table in detected_tables:
        table_index = int(table.get("table_index") or 0)
        mapping = mappings_by_table.get(table_index)
        if mapping is None or _is_unconfigured_default_mapping(mapping, table):
            pending_mapping_count += 1
        else:
            configured_mapping_count += 1

    high_mapping_suggestions = [
        suggestion
        for suggestion in mapping_suggestions
        if suggestion.get("confidence") == "high"
    ]
    high_mapping_pending_count = sum(
        1
        for suggestion in high_mapping_suggestions
        if (
            (mapping := mappings_by_table.get(int(suggestion.get("table_index") or 0))) is None
            or _is_unconfigured_default_mapping(
                mapping,
                next(
                    (
                        table
                        for table in detected_tables
                        if int(table.get("table_index") or 0) == int(suggestion.get("table_index") or 0)
                    ),
                    {},
                ),
            )
        )
    )

    slots_by_target = {
        (str(slot.get("target_kind") or "table_cell"), str(slot.get("target_locator") or "")): slot
        for slot in field_slots
    }
    high_field_slot_suggestions = [
        suggestion
        for suggestion in field_slot_suggestions
        if suggestion.get("confidence") == "high"
    ]
    high_field_slot_pending_count = sum(
        1
        for suggestion in high_field_slot_suggestions
        if (
            (slot := slots_by_target.get((
                str(suggestion.get("target_kind") or "table_cell"),
                str(suggestion.get("target_locator") or ""),
            ))) is None
            or _is_unconfigured_default_field_slot(slot)
        )
    )
    fixed_data_slot_count = sum(
        1
        for slot in field_slots
        if str(slot.get("value_source") or "") not in {"", "confirmed_text"}
    )
    confirmed_text_slot_count = sum(
        1
        for slot in field_slots
        if str(slot.get("value_source") or "") == "confirmed_text"
    )
    inline_token_suggestion_count = sum(
        1
        for suggestion in field_slot_suggestions
        if suggestion.get("target_kind") == "inline_token"
    )
    inline_token_slot_count = sum(
        1
        for slot in field_slots
        if slot.get("target_kind") == "inline_token"
    )
    return {
        "mapping": {
            "detected_table_count": len(detected_tables),
            "mapping_count": len(mappings),
            "configured_count": configured_mapping_count,
            "pending_count": pending_mapping_count,
            "high_confidence_suggestion_count": len(high_mapping_suggestions),
            "high_confidence_pending_count": high_mapping_pending_count,
        },
        "field_slots": {
            "configured_count": len(field_slots),
            "fixed_data_count": fixed_data_slot_count,
            "confirmed_text_count": confirmed_text_slot_count,
            "high_confidence_suggestion_count": len(high_field_slot_suggestions),
            "high_confidence_pending_count": high_field_slot_pending_count,
            "inline_token_suggestion_count": inline_token_suggestion_count,
            "inline_token_configured_count": inline_token_slot_count,
        },
        "outstanding_count": pending_mapping_count + high_field_slot_pending_count,
    }


def get_template_detail(template_id: str) -> dict[str, Any] | None:
    template = get_template(template_id)
    if template is None:
        return None
    detected_tables = _detected_tables_for_template(template)
    detected_text_targets = _detected_text_targets_for_template(template)
    template_path = Path(template["docx_path"])
    try:
        field_slot_suggestions = suggest_template_field_slots(template_path) if template_path.exists() else []
    except Exception:
        field_slot_suggestions = []
    mappings = list_template_mappings(template_id)
    field_slots = list_template_field_slots(template_id)
    mapping_suggestions = suggest_template_mappings(detected_tables)
    return {
        "template": template,
        "mappings": mappings,
        "field_slots": field_slots,
        "detected_tables": detected_tables,
        "mapping_suggestions": mapping_suggestions,
        "field_slot_suggestions": field_slot_suggestions,
        "detected_text_targets": detected_text_targets,
        "matching_profile": build_matching_profile(template, detected_tables),
        "configuration_readiness": _template_configuration_readiness(
            detected_tables=detected_tables,
            mappings=mappings,
            field_slots=field_slots,
            mapping_suggestions=mapping_suggestions,
            field_slot_suggestions=field_slot_suggestions,
        ),
    }


def get_template_recommendations(visit_type: str) -> dict[str, Any]:
    candidates = []
    for template in list_templates():
        detected_tables = _detected_tables_for_template(template)
        candidates.append(
            {
                "template": template,
                "detected_tables": detected_tables,
                "matching_profile": build_matching_profile(template, detected_tables),
            }
        )
    return recommend_templates(candidates, visit_type)


def update_template_visit_type_keywords(template_id: str, keywords: list[str]) -> dict[str, Any] | None:
    template = get_template(template_id)
    if template is None:
        return None
    if template["status"] not in {"draft", "rejected"}:
        raise ValueError("仅草稿或已退回模板可修改适用访视关键词；已启用模板请创建新草稿版本后配置")
    metadata = {**(template.get("metadata") or {})}
    metadata["visit_type_keywords"] = normalize_visit_type_keywords(keywords)
    update_template_control(template_id, {"metadata": metadata})
    return get_template_detail(template_id)


_TASK_COMPLETENESS_MODES = {"mapping_required", "all_mappings", "none"}
_FIELD_COMPLETENESS_MODES = {"slot_required", "all_confirmed_text_slots", "none"}


def _normalize_template_completeness_rules(rules: dict[str, Any]) -> dict[str, str]:
    task_mode = str(rules.get("task_mode") or "mapping_required").strip()
    field_mode = str(rules.get("field_mode") or "slot_required").strip()
    if task_mode not in _TASK_COMPLETENESS_MODES:
        raise ValueError("模板任务完整性模式无效")
    if field_mode not in _FIELD_COMPLETENESS_MODES:
        raise ValueError("模板字段完整性模式无效")
    return {"task_mode": task_mode, "field_mode": field_mode}


def update_template_completeness_rules(template_id: str, rules: dict[str, Any]) -> dict[str, Any] | None:
    template = get_template(template_id)
    if template is None:
        return None
    if template["status"] not in {"draft", "rejected"}:
        raise ValueError("仅草稿或已退回模板可修改完整性规则；已启用模板请创建新草稿版本后配置")
    metadata = {**(template.get("metadata") or {})}
    metadata["template_completeness_rules"] = _normalize_template_completeness_rules(rules)
    update_template_control(template_id, {"metadata": metadata})
    return get_template_detail(template_id)


def register_template(
    *,
    file_name: str,
    content: bytes,
    display_name: str,
    version: str,
    actor_name: str,
) -> dict[str, Any]:
    try:
        document = Document(BytesIO(content))
    except Exception as exc:
        raise ValueError("请上传可读取的 .docx Word 模板") from exc

    detected_tables = _detect_tables(document)
    TEMPLATE_UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    suffix = Path(file_name).suffix.lower() or ".docx"
    stored_path = TEMPLATE_UPLOAD_DIR / f"{uuid4().hex}{suffix}"
    stored_path.write_bytes(content)
    detected_text_targets = _detect_text_targets(document, stored_path)

    template_name = display_name.strip() or Path(file_name).stem or "未命名 Word 模板"
    matching_profile = build_matching_profile(
        {"name": template_name, "metadata": {"source_file_name": file_name}},
        detected_tables,
    )
    template = create_template(
        name=template_name,
        version=version.strip() or "V1.0",
        docx_path=str(stored_path.resolve()),
        table_count=len(detected_tables),
        metadata={
            "source": "uploaded_template",
            "export_profile": "generic_slots",
            "source_file_name": file_name,
            "uploaded_by": actor_name.strip() or "项目管理员",
            "detected_tables": detected_tables,
            "detected_text_targets": detected_text_targets,
            "matching_profile": matching_profile,
        },
    )
    replace_template_mappings(
        template["id"],
        [
            {
                "table_index": table["table_index"],
                "field_key": f"table_{table['table_index']}",
                "target_description": table["detected_label"],
                "required": False,
            }
            for table in detected_tables
        ],
    )
    if len(detected_tables) != 15:
        replace_template_field_slots(
            template["id"],
            [
                {
                    "table_index": table["table_index"],
                    "target_kind": "table_cell",
                    "label": f"{table['detected_label']} 填写位",
                    "field_key": f"table_{table['table_index']}",
                    "target_locator": table.get("suggested_target_locator") or "",
                    "value_source": "confirmed_text",
                    "required": False,
                }
                for table in detected_tables
                if table.get("suggested_target_locator")
            ],
        )
    return get_template_detail(template["id"]) or {}


def create_template_revision_draft(
    *,
    template_id: str,
    name: str,
    version: str,
    actor_name: str,
) -> dict[str, Any]:
    source_template = get_template(template_id)
    if source_template is None:
        raise ValueError("未找到 Word 模板")
    if source_template.get("status") not in {"active", "inactive"}:
        raise ValueError("仅已启用或已停用模板可以创建修订草稿")

    template_name = name.strip() or str(source_template.get("name") or "未命名 Word 模板")
    source_version = str(source_template.get("version") or "V1.0")
    template_version = version.strip() or f"{source_version}-R1"
    detected_tables = _detected_tables_for_template(source_template)
    metadata = {**(source_template.get("metadata") or {})}
    metadata.pop("activation_check", None)
    metadata["source"] = "template_revision_draft"
    metadata["revision_of"] = {
        "id": source_template["id"],
        "name": source_template.get("name", ""),
        "version": source_version,
        "status": source_template.get("status", ""),
        "created_by": actor_name.strip() or "项目管理员",
    }
    metadata["matching_profile"] = build_matching_profile(
        {"name": template_name, "metadata": metadata},
        detected_tables,
    )
    created = create_template(
        name=template_name,
        version=template_version,
        docx_path=str(source_template.get("docx_path") or ""),
        table_count=int(source_template.get("table_count") or 0),
        metadata=metadata,
    )
    replace_template_mappings(
        created["id"],
        [
            {
                "table_index": mapping["table_index"],
                "field_key": mapping["field_key"],
                "target_description": mapping["target_description"],
                "required": bool(mapping.get("required")),
            }
            for mapping in list_template_mappings(template_id)
        ],
    )
    replace_template_field_slots(
        created["id"],
        [
            {
                "table_index": slot["table_index"],
                "target_kind": slot.get("target_kind") or "table_cell",
                "label": slot.get("label") or "",
                "field_key": slot.get("field_key") or "",
                "target_locator": slot.get("target_locator") or "",
                "value_source": slot.get("value_source") or "confirmed_text",
                "required": bool(slot.get("required")),
            }
            for slot in list_template_field_slots(template_id)
        ],
    )
    return get_template_detail(created["id"]) or {}


def _reusable_field_slots_for_document(
    slots: list[dict[str, Any]],
    document: Document,
    detected_text_targets: list[dict[str, Any]],
) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    reusable: list[dict[str, Any]] = []
    removed: list[dict[str, Any]] = []
    text_target_keys = {
        (str(target.get("target_kind") or ""), str(target.get("target_locator") or ""))
        for target in detected_text_targets
    }
    for slot in slots:
        target_kind = str(slot.get("target_kind") or "table_cell")
        target_locator = str(slot.get("target_locator") or "")
        can_reuse = False
        if target_kind == "table_cell":
            try:
                target = parse_slot_target(target_kind, target_locator)
                table_index = int(target["table_index"])
                row_index = int(target["row_index"])
                column_index = int(target["column_index"])
                can_reuse = (
                    table_index <= len(document.tables)
                    and row_index <= len(document.tables[table_index - 1].rows)
                    and column_index <= len(document.tables[table_index - 1].rows[row_index - 1].cells)
                )
            except (KeyError, ValueError):
                can_reuse = False
        else:
            can_reuse = (target_kind, target_locator) in text_target_keys
        if can_reuse:
            reusable.append(slot)
        else:
            removed.append(slot)
    return reusable, removed


def replace_template_revision_document(
    *,
    template_id: str,
    file_name: str,
    content: bytes,
    actor_name: str,
) -> dict[str, Any]:
    template = get_template(template_id)
    if template is None:
        raise ValueError("未找到 Word 模板")
    if template.get("status") not in {"draft", "rejected"}:
        raise ValueError("仅修订草稿或已退回模板可替换 Word 文件")
    try:
        document = Document(BytesIO(content))
    except Exception as exc:
        raise ValueError("请上传可读取的 .docx Word 模板") from exc

    previous_detected_tables = _detected_tables_for_template(template)
    detected_tables = _detect_tables(document)
    TEMPLATE_UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    suffix = Path(file_name).suffix.lower() or ".docx"
    stored_path = TEMPLATE_UPLOAD_DIR / f"{uuid4().hex}{suffix}"
    stored_path.write_bytes(content)
    detected_text_targets = _detect_text_targets(document, stored_path)
    source_mappings = {int(item["table_index"]): item for item in list_template_mappings(template_id)}
    source_slots = list_template_field_slots(template_id)
    reusable_slots, removed_slots = _reusable_field_slots_for_document(source_slots, document, detected_text_targets)
    metadata = {**(template.get("metadata") or {})}
    metadata.pop("activation_check", None)
    metadata["source"] = "template_revision_document"
    metadata["export_profile"] = "generic_slots"
    metadata["source_file_name"] = file_name
    metadata["uploaded_by"] = actor_name.strip() or "项目管理员"
    metadata["detected_tables"] = detected_tables
    metadata["detected_text_targets"] = detected_text_targets
    metadata["document_replaced_from"] = {
        "source_file_name": str(template.get("metadata", {}).get("source_file_name") or ""),
        "table_count": int(template.get("table_count") or 0),
        "replaced_by": actor_name.strip() or "项目管理员",
    }
    metadata["document_replacement_summary"] = {
        "source_mapping_count": len(source_mappings),
        "reused_mapping_count": sum(1 for table in detected_tables if table["table_index"] in source_mappings),
        "source_field_slot_count": len(source_slots),
        "reused_field_slot_count": len(reusable_slots),
        "removed_field_slot_count": len(removed_slots),
        "removed_field_slot_labels": [str(slot.get("label") or "未命名填写位") for slot in removed_slots],
        **_describe_template_table_changes(previous_detected_tables, detected_tables),
    }
    metadata["matching_profile"] = build_matching_profile(
        {"name": template.get("name", ""), "metadata": metadata},
        detected_tables,
    )
    updated = update_template_document(
        template_id,
        docx_path=str(stored_path.resolve()),
        table_count=len(detected_tables),
        metadata=metadata,
    )
    if updated is None:
        raise ValueError("未找到 Word 模板")
    replace_template_mappings(
        template_id,
        [
            {
                "table_index": table["table_index"],
                "field_key": source_mappings.get(table["table_index"], {}).get("field_key") or f"table_{table['table_index']}",
                "target_description": source_mappings.get(table["table_index"], {}).get("target_description") or table["detected_label"],
                "required": bool(source_mappings.get(table["table_index"], {}).get("required")),
            }
            for table in detected_tables
        ],
    )
    replace_template_field_slots(template_id, reusable_slots)
    return get_template_detail(template_id) or {}


def export_template_configuration_package(template_id: str) -> dict[str, Any] | None:
    """Create a portable configuration package without including the source Word file."""
    detail = get_template_detail(template_id)
    if detail is None:
        return None
    template = detail["template"]
    metadata = template.get("metadata") or {}
    return {
        "package_type": "monitoring_mentor_template_configuration",
        "schema_version": "1.0",
        "exported_at": datetime.now().astimezone().isoformat(timespec="seconds"),
        "template": {
            "id": template["id"],
            "name": template.get("name", ""),
            "version": template.get("version", ""),
            "table_count": int(template.get("table_count") or 0),
            "source_file_name": str(metadata.get("source_file_name") or ""),
        },
        "configuration": {
            "mappings": [
                {
                    "table_index": int(mapping.get("table_index") or 0),
                    "field_key": str(mapping.get("field_key") or ""),
                    "target_description": str(mapping.get("target_description") or ""),
                    "required": bool(mapping.get("required")),
                }
                for mapping in detail["mappings"]
            ],
            "field_slots": [
                {
                    "table_index": int(slot.get("table_index") or 0),
                    "target_kind": str(slot.get("target_kind") or "table_cell"),
                    "label": str(slot.get("label") or ""),
                    "field_key": str(slot.get("field_key") or ""),
                    "target_locator": str(slot.get("target_locator") or ""),
                    "value_source": str(slot.get("value_source") or "confirmed_text"),
                    "required": bool(slot.get("required")),
                }
                for slot in detail["field_slots"]
            ],
            "visit_type_keywords": list(metadata.get("visit_type_keywords") or []),
            "template_completeness_rules": _normalize_template_completeness_rules(
                dict(metadata.get("template_completeness_rules") or {})
            ),
        },
    }


def _package_configuration(package: dict[str, Any]) -> dict[str, Any]:
    configuration = package.get("configuration", package)
    if not isinstance(configuration, dict):
        raise ValueError("模板配置包缺少 configuration 对象")
    return configuration


def _configuration_package_field_slots(configuration: dict[str, Any]) -> list[dict[str, Any]]:
    raw_slots = configuration.get("field_slots") or []
    if not isinstance(raw_slots, list):
        raise ValueError("模板配置包中的 field_slots 必须为数组")
    slots: list[dict[str, Any]] = []
    for raw_slot in raw_slots:
        if not isinstance(raw_slot, dict):
            continue
        try:
            table_index = int(raw_slot.get("table_index") or 0)
        except (TypeError, ValueError):
            table_index = 0
        slots.append(
            {
                "table_index": table_index,
                "target_kind": str(raw_slot.get("target_kind") or "table_cell").strip() or "table_cell",
                "label": str(raw_slot.get("label") or "").strip(),
                "field_key": str(raw_slot.get("field_key") or "").strip(),
                "target_locator": str(raw_slot.get("target_locator") or "").strip(),
                "value_source": str(raw_slot.get("value_source") or "confirmed_text").strip() or "confirmed_text",
                "required": bool(raw_slot.get("required")),
            }
        )
    return slots


def import_template_configuration_package(
    *,
    template_id: str,
    package: dict[str, Any],
    actor_name: str,
) -> dict[str, Any]:
    """Apply a reusable configuration package to an editable target template."""
    template = get_template(template_id)
    if template is None:
        raise ValueError("未找到 Word 模板")
    if template.get("status") not in {"draft", "rejected"}:
        raise ValueError("仅草稿或已退回模板可带入配置包")
    if not isinstance(package, dict):
        raise ValueError("请上传可读取的 JSON 模板配置包")

    configuration = _package_configuration(package)
    raw_mappings = configuration.get("mappings") or []
    if not isinstance(raw_mappings, list):
        raise ValueError("模板配置包中的 mappings 必须为数组")
    template_path = Path(str(template.get("docx_path") or ""))
    if not template_path.exists():
        raise ValueError("当前模板 Word 文件不可读取，无法带入配置包")
    try:
        document = Document(template_path)
    except Exception as exc:
        raise ValueError("当前模板 Word 文件不可读取，无法带入配置包") from exc

    detected_tables = _detected_tables_for_template(template)
    detected_text_targets = _detected_text_targets_for_template(template)
    source_mappings: dict[int, dict[str, Any]] = {}
    for raw_mapping in raw_mappings:
        if not isinstance(raw_mapping, dict):
            continue
        try:
            table_index = int(raw_mapping.get("table_index") or 0)
        except (TypeError, ValueError):
            continue
        if table_index > 0:
            source_mappings[table_index] = raw_mapping

    applied_mapping_count = 0
    target_mappings: list[dict[str, Any]] = []
    for table in detected_tables:
        table_index = int(table.get("table_index") or 0)
        source_mapping = source_mappings.get(table_index)
        if source_mapping is None:
            target_mappings.append(
                {
                    "table_index": table_index,
                    "field_key": f"table_{table_index}",
                    "target_description": str(table.get("detected_label") or f"表 {table_index}"),
                    "required": False,
                }
            )
            continue
        target_mappings.append(
            {
                "table_index": table_index,
                "field_key": str(source_mapping.get("field_key") or f"table_{table_index}").strip(),
                "target_description": str(source_mapping.get("target_description") or table.get("detected_label") or f"表 {table_index}").strip(),
                "required": bool(source_mapping.get("required")),
            }
        )
        applied_mapping_count += 1

    source_slots = _configuration_package_field_slots(configuration)
    reusable_slots, skipped_slots = _reusable_field_slots_for_document(source_slots, document, detected_text_targets)
    replace_template_mappings(template_id, target_mappings)
    replace_template_field_slots(template_id, reusable_slots)

    metadata = {**(template.get("metadata") or {})}
    keywords = configuration.get("visit_type_keywords")
    if isinstance(keywords, list):
        metadata["visit_type_keywords"] = normalize_visit_type_keywords([str(keyword) for keyword in keywords])
    rules = configuration.get("template_completeness_rules")
    if isinstance(rules, dict):
        metadata["template_completeness_rules"] = _normalize_template_completeness_rules(rules)
    source_template = package.get("template") if isinstance(package.get("template"), dict) else {}
    metadata["configuration_package_import"] = {
        "source_template_id": str(source_template.get("id") or ""),
        "source_template_name": str(source_template.get("name") or ""),
        "source_template_version": str(source_template.get("version") or ""),
        "schema_version": str(package.get("schema_version") or ""),
        "imported_by": actor_name.strip() or "项目管理员",
        "imported_at": datetime.now().astimezone().isoformat(timespec="seconds"),
        "source_mapping_count": len(raw_mappings),
        "applied_mapping_count": applied_mapping_count,
        "skipped_mapping_count": max(0, len(raw_mappings) - applied_mapping_count),
        "source_field_slot_count": len(source_slots),
        "applied_field_slot_count": len(reusable_slots),
        "skipped_field_slot_count": len(skipped_slots),
        "skipped_field_slot_labels": [str(slot.get("label") or "未命名填写位") for slot in skipped_slots],
    }
    metadata["matching_profile"] = build_matching_profile(
        {"name": template.get("name", ""), "metadata": metadata},
        detected_tables,
    )
    update_template_control(template_id, {"metadata": metadata})
    detail = get_template_detail(template_id) or {}
    return {
        "detail": detail,
        "source_template": {
            "name": str(source_template.get("name") or "未标注来源"),
            "version": str(source_template.get("version") or ""),
        },
        "source_mapping_count": len(raw_mappings),
        "applied_mapping_count": applied_mapping_count,
        "skipped_mapping_count": max(0, len(raw_mappings) - applied_mapping_count),
        "source_field_slot_count": len(source_slots),
        "applied_field_slot_count": len(reusable_slots),
        "skipped_field_slot_count": len(skipped_slots),
        "skipped_field_slot_labels": [str(slot.get("label") or "未命名填写位") for slot in skipped_slots],
    }


def import_high_confidence_template_field_slot_suggestions(
    *,
    template_id: str,
) -> dict[str, Any]:
    template = get_template(template_id)
    if template is None:
        raise ValueError("未找到 Word 模板")
    if template.get("status") not in {"draft", "rejected"}:
        raise ValueError("仅草稿或已退回模板可导入填写位建议")

    detail = get_template_detail(template_id) or {}
    suggestions = [
        suggestion
        for suggestion in list(detail.get("field_slot_suggestions") or [])
        if suggestion.get("confidence") == "high"
    ]
    existing_slots = list_template_field_slots(template_id)
    existing_by_target = {
        (str(slot.get("target_kind") or "table_cell"), str(slot.get("target_locator") or "")): slot
        for slot in existing_slots
    }
    default_slots_by_table = {
        int(slot.get("table_index") or 0): slot
        for slot in existing_slots
        if _is_unconfigured_default_field_slot(slot)
    }
    created_labels: list[str] = []
    adopted_labels: list[str] = []
    skipped_labels: list[str] = []

    for suggestion in suggestions:
        target_key = (str(suggestion.get("target_kind") or "table_cell"), str(suggestion.get("target_locator") or ""))
        existing = existing_by_target.get(target_key)
        if existing is None:
            default_slot = default_slots_by_table.get(int(suggestion.get("table_index") or 0))
            if default_slot is not None:
                update_template_field_slot(
                    template_id,
                    str(default_slot["id"]),
                    {
                        "target_kind": target_key[0],
                        "label": str(suggestion.get("label") or ""),
                        "field_key": str(suggestion.get("field_key") or ""),
                        "target_locator": target_key[1],
                        "value_source": str(suggestion.get("value_source") or "confirmed_text"),
                    },
                )
                default_slots_by_table.pop(int(suggestion.get("table_index") or 0), None)
                adopted_labels.append(str(suggestion.get("label") or "未命名填写位"))
                continue
            created = create_template_field_slot(
                template_id,
                {
                    "table_index": int(suggestion.get("table_index") or 0),
                    "target_kind": target_key[0],
                    "label": str(suggestion.get("label") or ""),
                    "field_key": str(suggestion.get("field_key") or ""),
                    "target_locator": target_key[1],
                    "value_source": str(suggestion.get("value_source") or "confirmed_text"),
                    "required": False,
                },
            )
            if created is not None:
                existing_by_target[target_key] = created
                created_labels.append(str(suggestion.get("label") or "未命名填写位"))
            continue
        if _is_unconfigured_default_field_slot(existing):
            update_template_field_slot(
                template_id,
                str(existing["id"]),
                {
                    "label": str(suggestion.get("label") or ""),
                    "field_key": str(suggestion.get("field_key") or ""),
                    "value_source": str(suggestion.get("value_source") or "confirmed_text"),
                },
            )
            adopted_labels.append(str(suggestion.get("label") or "未命名填写位"))
        else:
            skipped_labels.append(str(suggestion.get("label") or "未命名填写位"))

    return {
        "detail": get_template_detail(template_id) or {},
        "candidate_count": len(suggestions),
        "created_count": len(created_labels),
        "adopted_default_count": len(adopted_labels),
        "skipped_existing_count": len(skipped_labels),
        "created_labels": created_labels,
        "adopted_default_labels": adopted_labels,
        "skipped_existing_labels": skipped_labels,
    }


def import_high_confidence_template_mapping_suggestions(
    *,
    template_id: str,
) -> dict[str, Any]:
    template = get_template(template_id)
    if template is None:
        raise ValueError("未找到 Word 模板")
    if template.get("status") not in {"draft", "rejected"}:
        raise ValueError("仅草稿或已退回模板可导入映射建议")

    detected_tables = _detected_tables_for_template(template)
    tables_by_index = {int(table.get("table_index") or 0): table for table in detected_tables}
    suggestions = [
        suggestion
        for suggestion in suggest_template_mappings(detected_tables)
        if suggestion.get("confidence") == "high"
    ]
    mappings_by_table = {
        int(mapping.get("table_index") or 0): mapping
        for mapping in list_template_mappings(template_id)
    }
    adopted_labels: list[str] = []
    skipped_existing_labels: list[str] = []
    missing_mapping_labels: list[str] = []

    for suggestion in suggestions:
        table_index = int(suggestion.get("table_index") or 0)
        mapping = mappings_by_table.get(table_index)
        detected_table = tables_by_index.get(table_index)
        label = str(suggestion.get("target_description") or f"第 {table_index} 表")
        if mapping is None or detected_table is None:
            missing_mapping_labels.append(label)
            continue
        if not _is_unconfigured_default_mapping(mapping, detected_table):
            skipped_existing_labels.append(label)
            continue
        update_template_mapping(
            template_id,
            str(mapping["id"]),
            {
                "field_key": str(suggestion.get("field_key") or ""),
                "target_description": label,
            },
        )
        adopted_labels.append(label)

    return {
        "detail": get_template_detail(template_id) or {},
        "candidate_count": len(suggestions),
        "adopted_count": len(adopted_labels),
        "skipped_existing_count": len(skipped_existing_labels),
        "missing_mapping_count": len(missing_mapping_labels),
        "adopted_labels": adopted_labels,
        "skipped_existing_labels": skipped_existing_labels,
        "missing_mapping_labels": missing_mapping_labels,
    }


def create_template_field_slot_detail(template_id: str, payload: dict[str, Any]) -> dict[str, Any] | None:
    created = create_template_field_slot(template_id, payload)
    if created is None:
        return None
    return get_template_detail(template_id)


def update_template_field_slot_detail(template_id: str, slot_id: str, payload: dict[str, Any]) -> dict[str, Any] | None:
    updated = update_template_field_slot(template_id, slot_id, payload)
    if updated is None:
        return None
    return get_template_detail(template_id)


def delete_template_field_slot_detail(template_id: str, slot_id: str) -> dict[str, Any] | None:
    if not delete_template_field_slot(template_id, slot_id):
        return None
    return get_template_detail(template_id)
