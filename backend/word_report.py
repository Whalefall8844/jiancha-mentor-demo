from __future__ import annotations

from copy import deepcopy
from datetime import datetime
from pathlib import Path
from typing import Any
from uuid import uuid4

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

from .services.docx_structured_targets import apply_structured_target_values
from .services.template_slots import parse_slot_target, validate_slot_source


BACKEND_DIR = Path(__file__).resolve().parent
TEMPLATE_PATH = BACKEND_DIR / "templates" / "ua007_gt02_template.docx"
GENERATED_DIR = BACKEND_DIR / "generated"


def _short(value: Any, fallback: str = "不适用", limit: int = 220) -> str:
    text = str(value or "").strip()
    if not text:
        return fallback
    return text if len(text) <= limit else text[: limit - 1] + "…"


def _cell(table, row: int, col: int, value: Any, *, center: bool = False) -> None:
    cell = table.cell(row, col)
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(_short(value))
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(8)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        from docx.oxml import OxmlElement

        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def _mark(table, row: int, yes_col: int, no_col: int, na_col: int, choice: str, note: str = "") -> None:
    for col in (yes_col, no_col, na_col):
        _cell(table, row, col, "")
    selected = {"yes": yes_col, "no": no_col, "na": na_col}[choice]
    _cell(table, row, selected, "√", center=True)
    if note:
        _cell(table, row, na_col + 1, note)


def _confirmed_text(state: dict[str, Any], table_number: int) -> list[str]:
    return [
        _short(item.get("report_text") or item.get("value") or item.get("text"), limit=180)
        for item in state.get("confirmed_items", [])
        if item.get("target_table") == table_number
    ]


def _confirmed_subject(state: dict[str, Any], table_number: int, fallback: str) -> str:
    item = next(
        (entry for entry in state.get("confirmed_items", []) if entry.get("target_table") == table_number),
        None,
    )
    if item is None:
        return fallback
    return str(item.get("subject_code") or item.get("subject") or fallback)


def _activity_period(visit: dict[str, Any]) -> str:
    activity_end_date = str(visit.get("activity_end_date") or visit.get("visit_date") or "").strip()
    activity_start_date = str(visit.get("activity_start_date") or activity_end_date).strip() or activity_end_date
    if activity_start_date and activity_end_date and activity_start_date != activity_end_date:
        return f"{activity_start_date} 至 {activity_end_date}"
    return activity_end_date


def _activity_context_sentence(visit: dict[str, Any]) -> str:
    period = _activity_period(visit)
    if not period:
        return ""
    method = str(visit.get("visit_method") or "现场").strip() or "现场"
    location = str(visit.get("visit_location") or "").strip()
    contacts = str(visit.get("contact_persons") or "").strip()
    sentence = f"本次监查采用{method}方式，于{period}开展"
    if location:
        sentence += f"，通过{location}" if method == "远程" else f"，于{location}实施"
    if contacts:
        sentence += f"，并与{contacts}进行沟通"
    return f"{sentence}。"


def _derived_ae_records(state: dict[str, Any]) -> list[dict[str, Any]]:
    if state.get("ae_records"):
        return state["ae_records"]
    return [
        {
            "subject": finding.get("subject_code", ""),
            "description": finding.get("description", ""),
            "is_sae": finding.get("category") == "sae",
        }
        for finding in state.get("findings", [])
        if finding.get("category") in {"ae", "sae"}
    ]


def _derived_deviations(state: dict[str, Any]) -> list[dict[str, Any]]:
    if state.get("deviations"):
        return state["deviations"]
    return [
        {
            "subject": finding.get("subject_code", ""),
            "description": finding.get("description", ""),
        }
        for finding in state.get("findings", [])
        if finding.get("category") == "deviation"
    ]


def _summary(state: dict[str, Any]) -> str:
    lines = _confirmed_text(state, 3)
    if not lines:
        lines = [
            "本次为首例筛选监查访视。项目、中心及访视固定信息已载入；"
            "其余监查结论由 CRA 按实际工作底稿持续补录并确认。"
        ]
    activity_context = _activity_context_sentence(dict(state.get("visit") or {}))
    if activity_context:
        lines.append(activity_context)
    return " ".join(lines)


def _confirmed_text_by_field_key(state: dict[str, Any], field_key: str) -> str:
    seen: set[str] = set()
    values: list[str] = []
    for item in state.get("confirmed_items", []):
        if str(item.get("field_key") or "").strip() != field_key:
            continue
        text = str(item.get("report_text") or item.get("value") or "").strip()
        if text and text not in seen:
            seen.add(text)
            values.append(text)
    return "\n".join(values)


def _generic_slot_value(state: dict[str, Any], slot: dict[str, Any]) -> str:
    source = str(slot.get("value_source") or "").strip()
    field_key = str(slot.get("field_key") or "").strip()
    validate_slot_source(source, field_key)
    if source == "confirmed_text":
        return _confirmed_text_by_field_key(state, field_key)
    if source == "summary":
        return _summary(state)

    project = dict(state.get("project") or {})
    site = dict(state.get("site") or {})
    visit = dict(state.get("visit") or {})
    if source == "visit.activity_period":
        return _activity_period(visit)
    sources = {
        "project.study_name": project.get("study_name", ""),
        "project.study_id": project.get("study_id", ""),
        "project.sponsor": project.get("sponsor", ""),
        "project.approval_number": project.get("approval_number", ""),
        "project.sop_version": project.get("sop_version", ""),
        "site.site_name": site.get("site_name", ""),
        "site.pi_name": site.get("pi_name", ""),
        "site.protocol_version": site.get("protocol_version", ""),
        "site.icf_version": site.get("icf_version", ""),
        "site.ethics_date": site.get("ethics_date", ""),
        "visit.visit_method": visit.get("visit_method", ""),
        "visit.report_date": visit.get("report_date", ""),
        "visit.site_team": visit.get("site_team", ""),
        "visit.monitoring_team": visit.get("monitoring_team", ""),
        "visit.next_visit": visit.get("next_visit", ""),
    }
    return str(sources.get(source, "") or "").strip()


def _replace_cell_text_preserving_style(cell, value: str) -> None:
    paragraph = cell.paragraphs[0]
    original_rpr = deepcopy(paragraph.runs[0]._element.rPr) if paragraph.runs and paragraph.runs[0]._element.rPr is not None else None
    paragraph.clear()
    run = paragraph.add_run(_short(value, fallback="", limit=2000))
    if original_rpr is not None:
        run._element.insert(0, original_rpr)
    for extra_paragraph in cell.paragraphs[1:]:
        extra_paragraph.clear()


def _replace_paragraph_text_preserving_style(paragraph, value: str) -> None:
    original_rpr = deepcopy(paragraph.runs[0]._element.rPr) if paragraph.runs and paragraph.runs[0]._element.rPr is not None else None
    paragraph.clear()
    run = paragraph.add_run(_short(value, fallback="", limit=2000))
    if original_rpr is not None:
        run._element.insert(0, original_rpr)


def _paragraph_for_target(document: Document, target: dict[str, Any]):
    target_kind = str(target["target_kind"])
    paragraph_index = int(target["paragraph_index"])
    if target_kind == "body_paragraph" or target.get("region") == "body":
        return document.paragraphs[paragraph_index - 1]

    section_index = int(target["section_index"])
    section = document.sections[section_index - 1]
    is_header = target_kind == "header_paragraph" or target.get("region") == "header"
    paragraphs = section.header.paragraphs if is_header else section.footer.paragraphs
    return paragraphs[paragraph_index - 1]


def _replace_inline_token_preserving_style(paragraph, token: str, value: str) -> None:
    replacement = _short(value, fallback="", limit=2000)
    replaced_in_run = False
    for run in paragraph.runs:
        if token in run.text:
            run.text = run.text.replace(token, replacement)
            replaced_in_run = True
    if replaced_in_run:
        return

    paragraph_text = paragraph.text
    if token not in paragraph_text:
        raise ValueError("inline token was not found in the template paragraph")
    _replace_paragraph_text_preserving_style(paragraph, paragraph_text.replace(token, replacement))


def _slot_error(slot: dict[str, Any], detail: str) -> ValueError:
    label = str(slot.get("label") or "unnamed fill slot").strip()
    locator = str(slot.get("target_locator") or "missing locator").strip()
    return ValueError(f"report fill slot '{label}' ({locator}) {detail}")


def _structured_slot_entries(state: dict[str, Any], field_slots: list[dict[str, Any]]) -> list[dict[str, Any]]:
    entries: list[dict[str, Any]] = []
    for slot in field_slots:
        try:
            target = parse_slot_target(slot.get("target_kind"), str(slot.get("target_locator") or ""))
            if target["target_kind"] not in {"content_control", "bookmark", "merge_field"}:
                continue
            entries.append(
                {
                    "target": target,
                    "target_locator": str(slot.get("target_locator") or ""),
                    "label": str(slot.get("label") or ""),
                    "value": _short(_generic_slot_value(state, slot), fallback="", limit=2000),
                }
            )
        except (KeyError, ValueError) as exc:
            raise _slot_error(slot, f"could not resolve the structured target: {exc}") from exc
    return entries


def _fill_generic_template(document: Document, state: dict[str, Any], field_slots: list[dict[str, Any]]) -> None:
    for slot in field_slots:
        target = parse_slot_target(slot.get("target_kind"), str(slot.get("target_locator") or ""))
        if target["target_kind"] in {"content_control", "bookmark", "merge_field"}:
            continue
        if target["target_kind"] != "table_cell":
            try:
                paragraph = _paragraph_for_target(document, target)
                value = _generic_slot_value(state, slot)
                if target["target_kind"] == "inline_token":
                    _replace_inline_token_preserving_style(paragraph, str(target["token"]), value)
                else:
                    _replace_paragraph_text_preserving_style(paragraph, value)
            except (IndexError, KeyError, ValueError) as exc:
                raise _slot_error(slot, f"could not write the target: {exc}") from exc
            continue
        table_index = int(target["table_index"])
        row_index = int(target["row_index"])
        column_index = int(target["column_index"])
        try:
            table = document.tables[table_index - 1]
            cell = table.cell(row_index - 1, column_index - 1)
        except IndexError as exc:
            raise ValueError("报告填写位指向了模板中不存在的单元格") from exc
        _replace_cell_text_preserving_style(cell, _generic_slot_value(state, slot))


def _save_report(document: Document, state: dict[str, Any], revision_number: str | None) -> Path:
    GENERATED_DIR.mkdir(parents=True, exist_ok=True)
    project = dict(state.get("project") or {})
    visit = dict(state.get("visit") or {})
    visit_code = str(visit.get("code") or project.get("study_id") or "monitoring_report")
    safe_visit_code = "".join(character if character.isalnum() or character in {"-", "_"} else "_" for character in visit_code)
    version = revision_number or "Draft"
    filename = f"{safe_visit_code}_monitoring_report_{version}_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid4().hex[:6]}.docx"
    output_path = GENERATED_DIR / filename
    document.save(output_path)
    return output_path


def generate_report(state: dict[str, Any], revision_number: str | None = None) -> Path:
    configured_template = state.get("template", {}).get("docx_path")
    template_path = Path(configured_template) if configured_template else TEMPLATE_PATH
    if not template_path.is_absolute():
        template_path = BACKEND_DIR / template_path
    if not template_path.exists():
        raise FileNotFoundError(f"固定模板不存在：{template_path}")

    document = Document(template_path)
    field_slots = list(state.get("template_field_slots") or [])
    if field_slots:
        _fill_generic_template(document, state, field_slots)
        output_path = _save_report(document, state, revision_number)
        apply_structured_target_values(output_path, _structured_slot_entries(state, field_slots))
        return output_path
    template_metadata = dict(state.get("template", {}).get("metadata") or {})
    is_ua007_legacy_template = (
        template_metadata.get("export_profile") == "ua007_legacy_15_table"
        or template_metadata.get("source") == "fixed_demo_template"
    )
    if not is_ua007_legacy_template:
        raise ValueError("当前 Word 模板尚未配置报告填写位；请由项目管理员从识别建议中新增并保存填写位后再生成报告")
    if len(document.tables) != 15:
        raise ValueError(f"模板表格数为 {len(document.tables)}，预期应为 15")

    project = state["project"]
    visit = state["visit"]
    recruitment = state["recruitment"]
    tables = document.tables

    # 1. 研究项目固定信息
    t = tables[0]
    _cell(t, 0, 1, project.get("study_name"))
    _cell(t, 1, 1, project.get("pi_name"))
    _cell(t, 2, 1, project.get("site_name"))
    _cell(t, 3, 1, project.get("sponsor"))
    _cell(t, 4, 1, project.get("approval_number"))
    _cell(t, 5, 1, project.get("study_id"))
    _cell(t, 6, 1, project.get("sop_version"))

    # 2. 访视信息
    t = tables[1]
    activity_period = _activity_period(visit)
    visit_method = str(visit.get("visit_method") or "现场").strip() or "现场"
    _cell(t, 0, 2, f"{activity_period}（{visit_method}监查）" if activity_period else visit.get("visit_date"))
    _cell(t, 1, 2, visit.get("report_date"))
    _cell(t, 2, 2, visit.get("site_team"))
    _cell(t, 3, 2, visit.get("monitoring_team"))
    _cell(t, 4, 2, visit.get("next_visit"))

    # 3. 总体评价
    _cell(tables[2], 1, 1, _summary(state), center=False)

    # 4. 受试者招募
    t = tables[3]
    recruitment_rows = [
        ("screened", 12),
        ("screen_failed", 12),
        ("treated", 12),
        ("ae_dropout", 12),
        ("other_dropout", 12),
        ("completed_treatment", 12),
        ("follow_up", 12),
        ("follow_up_dropout", 12),
        ("completed_follow_up", 12),
    ]
    for row, (key, planned) in enumerate(recruitment_rows, start=1):
        _cell(t, row, 1, planned, center=True)
        _cell(t, row, 2, "NA", center=True)
        _cell(t, row, 3, recruitment.get(key, 0), center=True)

    # 5. 监查摘要
    t = tables[4]
    summary_notes = {
        1: "项目文件版本已载入，待现场逐项确认",
        2: "知情同意过程按 CRA 确认记录填报",
        3: "原始记录/CRF 情况见后续表格",
        4: "AE/SAE 情况见表 11",
        5: "文件审查情况见表 12",
        6: "偏离及违背情况见表 13",
        7: "研究药品核查见表 13",
        8: "其他监查事项见附加说明",
    }
    for row, note in summary_notes.items():
        _cell(t, row, 1, "是", center=True)
        _cell(t, row, 2, "是", center=True)
        _cell(t, row, 3, note)

    # 6. 法规文件
    t = tables[5]
    for row, note in ((1, project.get("protocol_version")), (2, project.get("ethics_date")), (3, project.get("icf_version"))):
        _mark(t, row, 3, 4, 5, "yes", note)
    versions = [
        (6, "方案", project.get("protocol_version")),
        (7, "修订版", "不适用"),
        (8, "知情同意书", project.get("icf_version")),
        (9, "病例报告表", "EDC/CRF 待项目配置"),
        (10, "研究者手册（如有）", "不适用"),
        (11, "伦理批准函", project.get("ethics_date")),
        (12, "其他（如有）", "不适用"),
    ]
    for row, item, version in versions:
        _cell(t, row, 1, item)
        _cell(t, row, 2, version)

    # 7. 知情同意
    t = tables[6]
    icf_notes = _confirmed_text(state, 7)
    for row in range(1, 7):
        _mark(t, row, 2, 3, 4, "yes", "已由 CRA 结合现场记录确认")
    _mark(t, 7, 2, 3, 4, "no", "；".join(icf_notes) if icf_notes else "暂未记录知情同意过程异常")

    # 8. 签署 ICF 的受试者
    t = tables[7]
    icf_list = _confirmed_text(state, 8)
    subject = _confirmed_subject(state, 8, "S-DEMO-001")
    _cell(t, 2, 1, subject)
    _cell(t, 2, 2, project.get("icf_version"))
    _cell(t, 2, 3, "；".join(icf_list) if icf_list else "演示数据，待 CRA 确认")

    # 9. CRF 审阅
    t = tables[8]
    crf_notes = _confirmed_text(state, 9)
    for row in range(1, 9):
        _mark(t, row, 2, 3, 4, "na", "；".join(crf_notes) if row == 1 and crf_notes else "本 Demo 暂未接入 EDC")

    # 10. CRF 清单及发现
    t = tables[9]
    _cell(t, 2, 1, "S-DEMO-001")
    for col in range(2, 7):
        _cell(t, 2, col, "NA", center=True)
    _cell(t, 11, 1, f"原始病历及病例报告表中的发现：已审核受试者人数：0；{_short('；'.join(_confirmed_text(state, 10)), '本 Demo 暂未接入 EDC', 160)}")
    _cell(t, 13, 1, "NA")
    _cell(t, 13, 2, "NA")
    _cell(t, 13, 3, "无 CRF 质询记录")
    _cell(t, 13, 5, "不适用")
    _cell(t, 13, 6, "CRA")

    # 11. AE/SAE
    t = tables[10]
    ae_records = _derived_ae_records(state)
    if ae_records:
        _mark(t, 2, 3, 4, 5, "yes")
        _mark(t, 3, 3, 4, 5, "yes" if any(item.get("is_sae") for item in ae_records) else "no")
        for row, item in enumerate(ae_records[:10], start=8):
            _cell(t, row, 1, item.get("subject", "未提供"))
            _cell(t, row, 2, item.get("description", "监查记录"))
            _cell(t, row, 3, "是" if item.get("is_sae") else "否", center=True)
            _cell(t, row, 7, "是", center=True)
    else:
        _mark(t, 2, 3, 4, 5, "no", "暂未记录 AE")
        _mark(t, 3, 3, 4, 5, "no", "暂未记录 SAE")
        _mark(t, 4, 3, 4, 5, "na", "不适用")
        _mark(t, 5, 3, 4, 5, "na", "不适用")
        _cell(t, 8, 1, "不适用")
        _cell(t, 8, 2, "本次访视暂未确认 AE/SAE 记录")
        _cell(t, 8, 3, "否", center=True)
        _cell(t, 8, 7, "不适用", center=True)

    # 12. 文件审核与存档
    t = tables[11]
    for row in range(1, 12):
        choice = "yes" if row in (4, 7, 8, 9, 11) else "na"
        note = "按本次访视记录确认" if choice == "yes" else "暂未到更新或审阅节点"
        _mark(t, row, 2, 3, 4, choice, note)

    # 13. 方案偏离、研究药品
    t = tables[12]
    deviations = _derived_deviations(state)
    for row in range(1, 11):
        choice = "yes" if row in (2, 4, 5, 9) else "na"
        _mark(t, row, 3, 4, 5, choice, "按 CRA 已确认记录填报")
    _mark(t, 10, 3, 4, 5, "yes" if deviations else "no", "详见下表" if deviations else "暂未确认方案偏离")
    if deviations:
        first = deviations[0]
        _cell(t, 13, 1, first.get("subject", "未提供"))
        _cell(t, 13, 2, first.get("description", "方案偏离记录"))
        _cell(t, 13, 4, "待 CRA 确认")
    else:
        _cell(t, 13, 1, "不适用")
        _cell(t, 13, 2, "暂未确认方案偏离或违背")
        _cell(t, 13, 4, "不适用")
    for row in range(15, 23):
        choice = "yes" if row in (15, 16, 18, 21) else "no"
        _mark(t, row, 3, 4, 5, choice, "研究药品事项待现场确认" if row == 21 else "演示数据")

    # 14. 附加说明及行动项
    t = tables[13]
    action_text = [item.get("description", "") for item in state.get("action_items", [])]
    findings = _confirmed_text(state, 14)
    _cell(t, 0, 0, "附加说明：本报告为本地 Demo 自动生成，最终内容须由 CRA 确认。")
    _cell(t, 1, 0, "过去检查结果情况：无历史 Demo 行动项，或待 CRA 补录。")
    _cell(t, 2, 0, "与相关人员讨论当前监查结果：" + _short("；".join(findings + action_text), "暂未记录需要升级的行动项", 500))

    # 15. 报告完成和审核信息
    t = tables[14]
    status_label = {"draft": "草稿", "submitted": "已提交待审核", "returned": "已退回待修订", "approved": "已批准"}.get(state.get("report_status"), "草稿")
    _cell(t, 0, 0, f"报告完成者: {visit.get('cra_name', '演示 CRA')}")
    _cell(t, 1, 0, "部门: 临床运营（演示）")
    _cell(t, 2, 0, f"日期: {visit.get('report_date', '')}")
    _cell(t, 4, 0, "审核者: PM/LM 审核人")
    _cell(t, 5, 0, f"审核状态: {status_label}")
    _cell(t, 6, 0, f"日期: {datetime.now().strftime('%Y-%m-%d')}")

    GENERATED_DIR.mkdir(parents=True, exist_ok=True)
    visit_code = str(visit.get("code") or project.get("study_id") or "monitoring_report")
    safe_visit_code = "".join(character if character.isalnum() or character in {"-", "_"} else "_" for character in visit_code)
    version = revision_number or "Draft"
    filename = f"{safe_visit_code}_monitoring_report_{version}_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid4().hex[:6]}.docx"
    output_path = GENERATED_DIR / filename
    document.save(output_path)
    return output_path
